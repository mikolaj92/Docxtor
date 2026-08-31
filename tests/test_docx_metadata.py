from __future__ import annotations

from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docxtor import inspect_docx_metadata, sanitize_docx_metadata


def _metadata_bytes(path: Path) -> bytes:
    with ZipFile(path) as archive:
        return b"\n".join(
            archive.read(name) for name in archive.namelist() if name.startswith("docProps/")
        )


def _metadata_values(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    with ZipFile(path) as archive:
        for name in ("docProps/core.xml", "docProps/app.xml"):
            if name not in archive.namelist():
                continue
            root = ET.fromstring(archive.read(name))
            for element in root.iter():
                if isinstance(element.text, str) and element.text.strip():
                    values[element.tag.rsplit("}", 1)[-1]] = element.text.strip()
    return values


def _add_custom_property(path: Path, value: str) -> None:
    custom_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
    xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <property name="PrivateOwner" pid="2" fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}">
    <vt:lpwstr>{value}</vt:lpwstr>
  </property>
</Properties>""".encode()
    with ZipFile(path, "a", compression=ZIP_DEFLATED) as archive:
        archive.writestr("docProps/custom.xml", custom_xml)


def test_sanitize_docx_metadata_removes_identity_and_custom_properties(tmp_path: Path) -> None:
    path = tmp_path / "metadata.docx"
    document = Document()
    document.add_paragraph("Visible text is not the only disclosure channel.")
    document.core_properties.author = "Jan Kowalski"
    document.core_properties.last_modified_by = "Anna Nowak"
    document.core_properties.title = "Poufna sprawa Jan Kowalski"
    document.core_properties.subject = "Anna Nowak"
    document.core_properties.keywords = "private"
    document.core_properties.description = "Secret description"
    document.core_properties.revision = 42
    document.save(path)
    _add_custom_property(path, "PrivateOwner")

    assert b"Jan Kowalski" in _metadata_bytes(path)
    assert not inspect_docx_metadata(path).clean

    sanitize_docx_metadata(path, review_author="Posejdon", review_initials="PD")

    assert _metadata_values(path) == {}
    inspection = inspect_docx_metadata(path)
    assert inspection.clean
    with ZipFile(path) as archive:
        assert "docProps/custom.xml" not in archive.namelist()
        package_rels = archive.read("_rels/.rels")
        assert b"custom-properties" not in package_rels
        content_types = archive.read("[Content_Types].xml")
        assert b"custom-properties" not in content_types
    assert b"Jan Kowalski" not in _metadata_bytes(path)
    assert b"Anna Nowak" not in _metadata_bytes(path)
    assert b"PrivateOwner" not in _metadata_bytes(path)


def test_sanitize_docx_metadata_is_safe_when_only_default_properties_exist(
    tmp_path: Path,
) -> None:
    path = tmp_path / "default-metadata.docx"
    document = Document()
    document.add_paragraph("No identity metadata")
    document.save(path)

    sanitize_docx_metadata(path, review_author="Posejdon", review_initials="PD")

    assert inspect_docx_metadata(path).clean
    assert Document(path).paragraphs[0].text == "No identity metadata"
