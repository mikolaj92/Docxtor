from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxtor import DocxDocument
from docxtor.docx_facts import (
    ChangeKind,
    FactsCoverage,
    TransformPolicy,
    compare_docx,
    docx_facts,
)
from docxtor.docx_package import (
    PackageEntry,
    PackageError,
    read_package_entries,
    write_package_atomically,
)


def _document(text: str = "hello") -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(stream)
    return stream.getvalue()


def _rewrite(tmp_path, data: bytes, changes: dict[str, bytes | None]) -> bytes:
    path = tmp_path / "changed.docx"
    entries = []
    seen = set()
    for entry in read_package_entries(data):
        seen.add(entry.name)
        replacement = changes.get(entry.name, entry.data)
        if replacement is not None:
            entries.append(PackageEntry(entry.name, replacement))
    for name, replacement in changes.items():
        if name not in seen and replacement is not None:
            entries.append(PackageEntry(name, replacement))
    write_package_atomically(path, entries)
    return path.read_bytes()


def test_snapshot_reports_parts_relationships_stories_and_orphan(tmp_path) -> None:
    data = _rewrite(tmp_path, _document(), {"custom/orphan.bin": b"opaque"})
    # Declare the opaque orphan so complete coverage is mechanically knowable.
    types = next(e.data for e in read_package_entries(data) if e.name == "[Content_Types].xml")
    types = types.replace(
        b"</Types>", b'<Default Extension="bin" ContentType="image/x-test"/></Types>'
    )
    data = _rewrite(tmp_path, data, {"[Content_Types].xml": types})

    snapshot = docx_facts(data)

    assert snapshot.coverage is FactsCoverage.COMPLETE
    assert "custom/orphan.bin" in snapshot.orphan_parts
    assert any(part.name == "custom/orphan.bin" and not part.reachable for part in snapshot.parts)
    assert snapshot.relationships
    assert snapshot.stories
    assert snapshot.paragraphs[0].coordinate.container_id == "body:p:0"


def test_notes_fields_bookmarks_and_hidden_are_typed(tmp_path) -> None:
    data = _document()
    document_xml = next(e.data for e in read_package_entries(data) if e.name == "word/document.xml")
    marker = b"<w:r><w:t>hello</w:t></w:r>"
    facts = (
        b'<w:bookmarkStart w:id="4" w:name="mark"/>'
        b'<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        b"<w:r><w:instrText>DATE</w:instrText></w:r>"
        b"<w:r><w:rPr><w:vanish/></w:rPr><w:t>secret</w:t></w:r>"
        b'<w:bookmarkEnd w:id="4"/>'
    )
    changed = _rewrite(
        tmp_path, data, {"word/document.xml": document_xml.replace(marker, marker + facts)}
    )

    snapshot = docx_facts(changed)

    assert {fact.value for fact in snapshot.fields} >= {"begin", "DATE"}
    assert {fact.value for fact in snapshot.bookmarks} >= {"mark", "4"}
    assert snapshot.hidden


def test_footnote_part_and_relationship_are_reported(tmp_path) -> None:
    data = _document()
    rels_name = "word/_rels/document.xml.rels"
    rels = next(e.data for e in read_package_entries(data) if e.name == rels_name)
    rel = (
        b'<Relationship Id="rIdFactsNote" '
        b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
        b'Target="footnotes.xml"/>'
    )
    rels = rels.replace(b"</Relationships>", rel + b"</Relationships>")
    types = next(e.data for e in read_package_entries(data) if e.name == "[Content_Types].xml")
    override = (
        b'<Override PartName="/word/footnotes.xml" '
        b'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
    )
    types = types.replace(b"</Types>", override + b"</Types>")
    notes = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:footnote w:id="1"><w:p><w:r><w:t>note</w:t></w:r></w:p></w:footnote>'
        b"</w:footnotes>"
    )
    changed = _rewrite(
        tmp_path, data, {rels_name: rels, "[Content_Types].xml": types, "word/footnotes.xml": notes}
    )

    snapshot = docx_facts(changed)

    assert any(note.value == "1" and note.target == "note" for note in snapshot.notes)
    assert "word/footnotes.xml" in snapshot.reachable_parts


def test_malformed_xml_and_missing_relationship_target_fail_closed(tmp_path) -> None:
    data = _document()
    with pytest.raises(PackageError):
        docx_facts(_rewrite(tmp_path, data, {"word/document.xml": b"<broken>"}))

    rels_name = "word/_rels/document.xml.rels"
    rels = next(e.data for e in read_package_entries(data) if e.name == rels_name)
    rel = b'<Relationship Id="rMissing" Type="urn:test" Target="missing.xml"/>'
    with pytest.raises(PackageError, match="missing"):
        docx_facts(
            _rewrite(
                tmp_path,
                data,
                {rels_name: rels.replace(b"</Relationships>", rel + b"</Relationships>")},
            )
        )


def test_compare_detects_lost_container_surface_part_and_policy(tmp_path) -> None:
    before = _document("before")
    after = _document("after")

    comparison = compare_docx(before, after)

    assert not comparison.allowed
    assert any(change.kind is ChangeKind.CHANGED for change in comparison.container_changes)
    assert comparison.surface_changes
    assert comparison.part_changes

    policy = TransformPolicy(
        allowed_part_changes=frozenset({ChangeKind.CHANGED}),
        allowed_surface_changes=frozenset({ChangeKind.CHANGED}),
        allowed_container_changes=frozenset({ChangeKind.CHANGED}),
        allowed_fact_categories=frozenset({"*"}),
    )
    assert compare_docx(before, after, policy).allowed

    empty = _document("")
    lost = compare_docx(before, empty, TransformPolicy.allow_all())
    assert any(change.kind is ChangeKind.CHANGED for change in lost.container_changes)


def test_feature_facts_bind_hidden_text_to_paragraph_and_resolve_link(tmp_path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("same ")
    hidden = paragraph.add_run("secret")
    hidden.font.hidden = True
    document.add_paragraph("same secret")
    stream = BytesIO()
    document.save(stream)

    snapshot = docx_facts(stream.getvalue())

    fact = next(item for item in snapshot.hidden if item.target == "secret")
    assert fact.container_id == "body:p:0"
    paragraph_fact = next(item for item in snapshot.paragraphs if item.container_id == "body:p:0")
    assert paragraph_fact.xml_path
    assert paragraph_fact.part_name == "word/document.xml"


def test_all_story_and_safety_facts_are_consumer_ready(tmp_path) -> None:
    document = Document()
    paragraph = document.add_paragraph("Visible ")
    hidden = paragraph.add_run("not-hidden")
    hidden.font.hidden = False
    toc = document.add_paragraph("Entry\t2")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "TOC1")
    toc._p.get_or_add_pPr().insert(0, style)
    stream = BytesIO()
    document.save(stream)
    data = stream.getvalue()
    notes = (
        b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:footnote w:id="-1" w:type="separator"><w:p/></w:footnote>'
        b'<w:footnote w:id="2"><w:p><w:r><w:t>note</w:t></w:r></w:p></w:footnote>'
        b"</w:footnotes>"
    )
    changed = _rewrite(tmp_path, data, {"word/footnotes.xml": notes})

    snapshot = docx_facts(changed)

    assert not snapshot.hidden
    assert any(item.style_id == "TOC1" for item in snapshot.paragraphs)
    assert any(
        item.container_id == "footnote:2:p:0" and item.text == "note"
        for item in snapshot.paragraphs
    )


def test_unreadable_properties_are_typed_without_consumer_xml_parse(tmp_path) -> None:
    data = _document()
    changed_stream = BytesIO()
    with (
        ZipFile(BytesIO(data)) as source,
        ZipFile(changed_stream, "w", compression=ZIP_DEFLATED) as target,
    ):
        for info in source.infolist():
            target.writestr(
                info,
                b"<not-xml" if info.filename == "docProps/app.xml" else source.read(info),
            )

    snapshot = docx_facts(changed_stream.getvalue())

    assert snapshot.coverage is FactsCoverage.INCOMPLETE
    assert snapshot.unreadable_parts[0].part_name == "docProps/app.xml"


def test_document_facts_use_original_package_including_orphan_parts(tmp_path) -> None:
    data = _document()
    changed = _rewrite(
        tmp_path,
        data,
        {"word/embeddings/Microsoft_Excel_Worksheet.xlsx": b"PK\x03\x04opaque"},
    )
    document = DocxDocument.open_bytes(changed)

    snapshot = document.facts()

    assert any(
        part.name == "word/embeddings/Microsoft_Excel_Worksheet.xlsx" for part in snapshot.parts
    )
