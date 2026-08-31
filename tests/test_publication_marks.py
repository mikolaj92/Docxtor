from io import BytesIO

from docx import Document

from docxtor import (
    BodyAppendix,
    DocumentMark,
    append_body_appendix,
    has_body_appendix,
    has_document_mark,
    remove_body_appendix,
    stamp_document_mark,
)


def _data() -> bytes:
    document = Document()
    document.add_paragraph("original")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_mark_and_appendix_are_typed_physical_operations() -> None:
    mark = DocumentMark("final", "Final", ("one", "two"), "Final artifact")
    marked = stamp_document_mark(_data(), mark)
    assert has_document_mark(marked, mark)
    appendix = BodyAppendix("Review trail", ("Intro", "1. [A] note"))
    appended = append_body_appendix(marked, appendix)
    assert has_body_appendix(appended, heading=appendix.heading)
    cleaned, changed = remove_body_appendix(appended, heading=appendix.heading)
    assert changed and not has_body_appendix(cleaned, heading=appendix.heading)
    assert Document(BytesIO(cleaned)).paragraphs[0].text == "original"
