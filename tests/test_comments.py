from __future__ import annotations

from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document as PyDocxDocument

from docxtor import AddressableComment, DocxDocument, SegmentReplacement, UnsupportedRevisionError

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
_REL = "{http://schemas.openxmlformats.org/package/2006/relationships}"
_CT = "{http://schemas.openxmlformats.org/package/2006/content-types}"

_COMMENTS_EXTENDED_TYPE = "application/vnd.ms-word.commentsExtended+xml"
_COMMENTS_IDS_TYPE = "application/vnd.ms-word.commentsIds+xml"
_PEOPLE_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.people+xml"
_COMMENTS_EXTENDED_REL = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
_COMMENTS_IDS_REL = "http://schemas.microsoft.com/office/2016/relationships/commentsIds"
_PEOPLE_REL = "http://schemas.microsoft.com/office/2011/relationships/people"
_FOOTNOTES_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
_ENDNOTES_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
_COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
_FOOTNOTES_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
_ENDNOTES_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"


def _plain_comment_docx(
    path: Path,
    *,
    body: str = "Clause text.",
    note: str = "check this clause",
    author: str = "Ann Reviewer",
    initials: str = "AR",
) -> Path:
    source = PyDocxDocument()
    paragraph = source.add_paragraph(body)
    source.add_comment(
        runs=paragraph.runs[0],
        text=note,
        author=author,
        initials=initials,
    )
    source.save(str(path))
    return path


def _multi_run_comment_docx(path: Path) -> Path:
    source = PyDocxDocument()
    paragraph = source.add_paragraph("Anchored body.")
    comment = source.add_comment(
        runs=paragraph.runs[0],
        text="Hello",
        author="Ann Reviewer",
        initials="AR",
    )
    comment.paragraphs[0].add_run(" world")
    source.save(str(path))
    return path


def _zip_replace(path: Path, updates: dict[str, bytes]) -> None:
    with ZipFile(path) as bundle:
        original = {name: bundle.read(name) for name in bundle.namelist()}
    original.update(updates)
    with ZipFile(path, mode="w", compression=ZIP_DEFLATED) as bundle:
        for name, data in original.items():
            bundle.writestr(name, data)


def _inject_reply_and_sidecars(path: Path) -> Path:
    with ZipFile(path) as bundle:
        comments_xml = bundle.read("word/comments.xml")
        document_rels = bundle.read("word/_rels/document.xml.rels")
        content_types = bundle.read("[Content_Types].xml")

    comments_root = ElementTree.fromstring(comments_xml)
    comment_paras = comments_root.findall(f"{_W}comment/{_W}p")
    assert comment_paras, "expected a parent comment paragraph"
    comment_paras[0].set(f"{_W14}paraId", "AAAA0001")

    reply = ElementTree.SubElement(comments_root, f"{_W}comment")
    reply.set(f"{_W}id", "1")
    reply.set(f"{_W}author", "Bob Reviewer")
    reply.set(f"{_W}initials", "BR")
    reply.set(f"{_W}date", "2024-03-04T10:11:12Z")
    reply_p = ElementTree.SubElement(reply, f"{_W}p")
    reply_p.set(f"{_W14}paraId", "BBBB0002")
    reply_r = ElementTree.SubElement(reply_p, f"{_W}r")
    reply_t = ElementTree.SubElement(reply_r, f"{_W}t")
    reply_t.text = "agreed, verify"

    rels_root = ElementTree.fromstring(document_rels)
    for rel_id, rel_type, target in (
        ("rIdCommentEx", _COMMENTS_EXTENDED_REL, "commentsExtended.xml"),
        ("rIdCommentIds", _COMMENTS_IDS_REL, "commentsIds.xml"),
        ("rIdPeople", _PEOPLE_REL, "people.xml"),
    ):
        rel = ElementTree.SubElement(rels_root, f"{_REL}Relationship")
        rel.set("Id", rel_id)
        rel.set("Type", rel_type)
        rel.set("Target", target)

    types_root = ElementTree.fromstring(content_types)
    for part_name, content_type in (
        ("/word/commentsExtended.xml", _COMMENTS_EXTENDED_TYPE),
        ("/word/commentsIds.xml", _COMMENTS_IDS_TYPE),
        ("/word/people.xml", _PEOPLE_TYPE),
    ):
        override = ElementTree.SubElement(types_root, f"{_CT}Override")
        override.set("PartName", part_name)
        override.set("ContentType", content_type)

    extended = (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        f'<w15:commentsEx xmlns:w15="{_W15}">'
        '<w15:commentEx w15:paraId="AAAA0001" w15:done="0"/>'
        '<w15:commentEx w15:paraId="BBBB0002" w15:paraIdParent="AAAA0001" w15:done="0"/>'
        "</w15:commentsEx>"
    ).encode()
    comment_ids = (
        b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        b'<w16cid:commentsIds xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid">'
        b'<w16cid:commentId w16cid:paraId="AAAA0001" w16cid:durableId="11111111"/>'
        b'<w16cid:commentId w16cid:paraId="BBBB0002" w16cid:durableId="22222222"/>'
        b"</w16cid:commentsIds>"
    )
    people = (
        b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        b'<w15:people xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
        b'<w15:person w15:author="Ann Reviewer"/>'
        b"</w15:people>"
    )

    _zip_replace(
        path,
        {
            "word/comments.xml": ElementTree.tostring(
                comments_root, encoding="UTF-8", xml_declaration=True
            ),
            "word/_rels/document.xml.rels": ElementTree.tostring(
                rels_root, encoding="UTF-8", xml_declaration=True
            ),
            "[Content_Types].xml": ElementTree.tostring(
                types_root, encoding="UTF-8", xml_declaration=True
            ),
            "word/commentsExtended.xml": extended,
            "word/commentsIds.xml": comment_ids,
            "word/people.xml": people,
        },
    )
    return path


def _write_parts(
    path: Path,
    document_xml: str,
    *,
    extras: dict[str, str] | None = None,
    document_rels: str | None = None,
    overrides: str = "",
) -> None:
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    office_rel = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
    )
    main_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
    rels_ct = "application/vnd.openxmlformats-package.relationships+xml"
    extras = extras or {}
    with ZipFile(path, mode="w", compression=ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<Types xmlns="{ct_ns}">'
                f'<Default Extension="rels" ContentType="{rels_ct}"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                f'<Override PartName="/word/document.xml" ContentType="{main_ct}"/>'
                f"{overrides}"
                "</Types>"
            ),
        )
        archive.writestr(
            "_rels/.rels",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<Relationships xmlns="{rel_ns}">'
                f'<Relationship Id="rId1" Type="{office_rel}" Target="word/document.xml"/>'
                "</Relationships>"
            ),
        )
        archive.writestr("word/document.xml", document_xml)
        if document_rels is not None:
            archive.writestr("word/_rels/document.xml.rels", document_rels)
        for name, xml in extras.items():
            archive.writestr(name, xml)


def _comment_markers(path: Path) -> list[tuple[str, str]]:
    root = ElementTree.fromstring(ZipFile(path).read("word/document.xml"))
    markers: list[tuple[str, str]] = []
    for element in root.iter():
        local = element.tag.rsplit("}", 1)[-1]
        if local in {"commentRangeStart", "commentRangeEnd", "commentReference"}:
            markers.append((local, element.get(f"{_W}id") or ""))
    return markers


def _comment_authors(path: Path) -> dict[str, str]:
    root = ElementTree.fromstring(ZipFile(path).read("word/comments.xml"))
    return {
        (comment.get(f"{_W}id") or ""): (comment.get(f"{_W}author") or "")
        for comment in root.findall(f"{_W}comment")
    }


def _comment_tree(path: Path) -> ElementTree.Element:
    return ElementTree.fromstring(ZipFile(path).read("word/comments.xml"))


def _comment_attr(path: Path, comment_id: str, attr: str) -> str | None:
    for comment in _comment_tree(path).iter():
        if comment.tag != f"{_W}comment":
            continue
        if comment.get(f"{_W}id") == comment_id:
            return comment.get(attr)
    return None


def test_populated_comments_expose_identity_offsets_anchor_and_parent(
    tmp_path: Path,
) -> None:
    path = _inject_reply_and_sidecars(_plain_comment_docx(tmp_path / "thread.docx"))

    doc = DocxDocument.open(path)

    assert [segment.container_id for segment in doc.segments if segment.container_id][:1] == [
        "body:p:0"
    ]
    comment_segments = [
        segment for segment in doc.segments if (segment.container_id or "").startswith("comment:")
    ]
    assert [segment.container_id for segment in comment_segments] == [
        "comment:0:p:0",
        "comment:1:p:0",
    ]
    assert comment_segments[0].text == "check this clause"
    assert comment_segments[1].text == "agreed, verify"

    comments = list(doc.comments)
    assert comments == [
        AddressableComment(
            comment_id="0",
            container_id="comment:0:p:0",
            text="check this clause",
            author="Ann Reviewer",
            initials="AR",
            locator="body:p:0",
            anchor_text="Clause text.",
            parent_id=None,
            date=comments[0].date,
        ),
        AddressableComment(
            comment_id="1",
            container_id="comment:1:p:0",
            text="agreed, verify",
            author="Bob Reviewer",
            initials="BR",
            locator=None,
            anchor_text="",
            parent_id="0",
            date="2024-03-04T10:11:12Z",
        ),
    ]
    parent_span = next(span for span in doc.spans if span.container_id == "comment:0:p:0")
    assert parent_span.text == "check this clause"
    assert parent_span.start_offset == 0
    assert parent_span.end_offset == len("check this clause")
    assert parent_span.span_id == "comment:0:p:0:span:0"


def test_comment_text_replacement_preserves_ids_anchors_metadata_and_sidecars(
    tmp_path: Path,
) -> None:
    path = _inject_reply_and_sidecars(_plain_comment_docx(tmp_path / "edit.docx"))
    out = tmp_path / "edit-out.docx"
    original_markers = _comment_markers(path)
    original_authors = _comment_authors(path)

    doc = DocxDocument.open(path)
    body_text = doc.texts[0]
    doc.apply_replacements(
        [
            SegmentReplacement(container_id="comment:0:p:0", text="redacted note"),
            SegmentReplacement(
                span_id="comment:1:p:0:span:0",
                text="X",
                start_offset=0,
                end_offset=6,
            ),
        ],
        strict=True,
    )
    doc.save_docx(out)

    reopened = DocxDocument.open(out)
    assert reopened.texts[0] == body_text
    by_id = {comment.comment_id: comment for comment in reopened.comments}
    assert by_id["0"].text == "redacted note"
    assert by_id["0"].author == "Ann Reviewer"
    assert by_id["0"].initials == "AR"
    assert by_id["0"].locator == "body:p:0"
    assert by_id["0"].anchor_text == "Clause text."
    assert by_id["1"].text == "X, verify"
    assert by_id["1"].parent_id == "0"
    assert by_id["1"].author == "Bob Reviewer"
    assert _comment_markers(out) == original_markers
    assert _comment_authors(out) == original_authors

    names = set(ZipFile(out).namelist())
    assert "word/commentsExtended.xml" in names
    assert "word/commentsIds.xml" in names
    assert "word/people.xml" in names
    comments_xml = ZipFile(out).read("word/comments.xml").decode("utf-8")
    assert "check this clause" not in comments_xml
    assert _comment_attr(out, "0", f"{_W}id") == "0"
    assert _comment_attr(out, "0", f"{_W}author") == "Ann Reviewer"
    parent_para = next(
        paragraph
        for comment in _comment_tree(out).findall(f"{_W}comment")
        if comment.get(f"{_W}id") == "0"
        for paragraph in comment.findall(f"{_W}p")
    )
    assert parent_para.get(f"{_W14}paraId") == "AAAA0001"
    extended = ZipFile(out).read("word/commentsExtended.xml")
    assert b"BBBB0002" in extended and b"AAAA0001" in extended
    assert b'paraIdParent="AAAA0001"' in extended
    assert b"11111111" in ZipFile(out).read("word/commentsIds.xml")
    assert b"Ann Reviewer" in ZipFile(out).read("word/people.xml")
    rels = ZipFile(out).read("word/_rels/document.xml.rels").decode("utf-8")
    assert "commentsExtended.xml" in rels
    assert "commentsIds.xml" in rels
    assert "people.xml" in rels


def test_multiple_comment_runs_are_one_addressable_segment(tmp_path: Path) -> None:
    path = _multi_run_comment_docx(tmp_path / "runs.docx")
    out = tmp_path / "runs-out.docx"

    doc = DocxDocument.open(path)
    comment = doc.comments[0]
    assert comment.text == "Hello world"
    assert comment.container_id == "comment:0:p:0"
    assert comment.locator == "body:p:0"
    assert comment.anchor_text == "Anchored body."

    doc.apply_replacements(
        [
            SegmentReplacement(
                container_id="comment:0:p:0",
                text="Hi",
                start_offset=0,
                end_offset=5,
            )
        ],
        strict=True,
    )
    doc.save_docx(out)
    reopened = DocxDocument.open(out)
    assert reopened.comments[0].text == "Hi world"
    assert reopened.comments[0].comment_id == "0"
    assert "Hello" not in ZipFile(out).read("word/comments.xml").decode("utf-8")


def test_empty_comments_and_separator_notes_are_not_user_segments(
    tmp_path: Path,
) -> None:
    empty_comments = tmp_path / "empty-comments.docx"
    _write_parts(
        empty_comments,
        """
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    <w:p><w:r><w:t>Visible.</w:t></w:r></w:p>
  </w:body>
</w:document>
""".strip(),
        document_rels=(
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            f'<Relationship Id="rIdComments" Type="{_COMMENTS_REL}" Target="comments.xml"/>'
            f'<Relationship Id="rIdFootnotes" Type="{_FOOTNOTES_REL}" Target="footnotes.xml"/>'
            "</Relationships>"
        ),
        extras={
            "word/comments.xml": (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            ),
            "word/footnotes.xml": (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1">'
                "<w:p><w:r><w:separator/></w:r></w:p>"
                "</w:footnote>"
                '<w:footnote w:type="continuationSeparator" w:id="0">'
                "<w:p><w:r><w:continuationSeparator/></w:r></w:p>"
                "</w:footnote>"
                "</w:footnotes>"
            ),
        },
        overrides=(
            '<Override PartName="/word/comments.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
            f'<Override PartName="/word/footnotes.xml" ContentType="{_FOOTNOTES_TYPE}"/>'
        ),
    )

    doc = DocxDocument.open(empty_comments)
    assert list(doc.comments) == []
    assert [segment.container_id for segment in doc.segments] == ["body:p:0"]
    assert doc.texts == ["Visible."]
    assert not any((segment.container_id or "").startswith("comment:") for segment in doc.segments)
    assert not any((segment.container_id or "").startswith("footnote:") for segment in doc.segments)


def test_user_footnotes_and_endnotes_are_addressable_and_round_trip(tmp_path: Path) -> None:
    source_path = tmp_path / "notes.docx"
    output_path = tmp_path / "notes-out.docx"
    _write_parts(
        source_path,
        """
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p><w:r><w:t>Visible body.</w:t></w:r></w:p></w:body>
</w:document>
""".strip(),
        document_rels=(
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            f'<Relationship Id="rIdFootnotes" Type="{_FOOTNOTES_REL}" Target="footnotes.xml"/>'
            f'<Relationship Id="rIdEndnotes" Type="{_ENDNOTES_REL}" Target="endnotes.xml"/>'
            "</Relationships>"
        ),
        extras={
            "word/footnotes.xml": (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:type="separator" w:id="-1">'
                '<w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                '<w:footnote w:id="2"><w:p><w:r><w:t>Private footnote</w:t></w:r></w:p>'
                '<w:p><w:r><w:t>Second paragraph</w:t></w:r></w:p></w:footnote>'
                '</w:footnotes>'
            ),
            "word/endnotes.xml": (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:endnote w:type="continuationSeparator" w:id="0">'
                '<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>'
                '<w:endnote w:id="7"><w:p><w:r><w:t>Private endnote</w:t></w:r></w:p></w:endnote>'
                '</w:endnotes>'
            ),
        },
        overrides=(
            f'<Override PartName="/word/footnotes.xml" ContentType="{_FOOTNOTES_TYPE}"/>'
            f'<Override PartName="/word/endnotes.xml" ContentType="{_ENDNOTES_TYPE}"/>'
        ),
    )

    document = DocxDocument.open(source_path)
    assert [segment.container_id for segment in document.segments] == [
        "body:p:0",
        "footnote:2:p:0",
        "footnote:2:p:1",
        "endnote:7:p:0",
    ]
    document.apply_replacements(
        [
            SegmentReplacement(container_id="footnote:2:p:0", text="Masked footnote"),
            SegmentReplacement(container_id="endnote:7:p:0", text="Masked endnote"),
        ],
        strict=True,
    )
    document.save_docx(output_path)

    reopened = DocxDocument.open(output_path)
    assert reopened.texts == [
        "Visible body.",
        "Masked footnote",
        "Second paragraph",
        "Masked endnote",
    ]
    with ZipFile(output_path) as bundle:
        assert b"Private footnote" not in bundle.read("word/footnotes.xml")
        assert b"Private endnote" not in bundle.read("word/endnotes.xml")
        assert b"separator" in bundle.read("word/footnotes.xml")
        assert b"continuationSeparator" in bundle.read("word/endnotes.xml")


def test_documents_without_comments_and_untouched_comments_do_not_drift(
    tmp_path: Path,
) -> None:
    plain = tmp_path / "plain.docx"
    plain_out = tmp_path / "plain-out.docx"
    source = PyDocxDocument()
    source.add_paragraph("Body only")
    source.save(str(plain))
    assert "word/comments.xml" not in ZipFile(plain).namelist()

    DocxDocument.open(plain).save_docx(plain_out)
    assert "word/comments.xml" not in ZipFile(plain_out).namelist()
    assert DocxDocument.open(plain_out).comments == ()

    threaded = _inject_reply_and_sidecars(_plain_comment_docx(tmp_path / "untouched.docx"))
    untouched_out = tmp_path / "untouched-out.docx"
    original_comments = ZipFile(threaded).read("word/comments.xml")
    original_extended = ZipFile(threaded).read("word/commentsExtended.xml")
    original_ids = ZipFile(threaded).read("word/commentsIds.xml")
    original_people = ZipFile(threaded).read("word/people.xml")
    original_markers = _comment_markers(threaded)

    doc = DocxDocument.open(threaded)
    assert [comment.comment_id for comment in doc.comments] == ["0", "1"]
    doc.save_docx(untouched_out)

    assert _comment_markers(untouched_out) == original_markers
    assert ZipFile(untouched_out).read("word/commentsExtended.xml") == original_extended
    assert ZipFile(untouched_out).read("word/commentsIds.xml") == original_ids
    assert ZipFile(untouched_out).read("word/people.xml") == original_people
    reopened = DocxDocument.open(untouched_out)
    assert [comment.text for comment in reopened.comments] == [
        "check this clause",
        "agreed, verify",
    ]
    assert [comment.parent_id for comment in reopened.comments] == [None, "0"]
    saved_comments = ZipFile(untouched_out).read("word/comments.xml").decode("utf-8")
    assert "check this clause" in saved_comments
    assert _comment_attr(untouched_out, "0", f"{_W}author") == "Ann Reviewer"
    parent_para = next(
        paragraph
        for comment in _comment_tree(untouched_out).findall(f"{_W}comment")
        if comment.get(f"{_W}id") == "0"
        for paragraph in comment.findall(f"{_W}p")
    )
    assert parent_para.get(f"{_W14}paraId") == "AAAA0001"
    assert original_comments  # source existed; rewrite may not be byte-identical


def test_unknown_comment_target_fails_closed_before_partial_write(tmp_path: Path) -> None:
    path = _plain_comment_docx(tmp_path / "fail-closed.docx")
    doc = DocxDocument.open(path)
    original_texts = list(doc.texts)
    original_comments = [comment.text for comment in doc.comments]

    with pytest.raises(ValueError, match="comment:999"):
        doc.apply_replacements(
            [
                SegmentReplacement(container_id="body:p:0", text="CHANGED"),
                SegmentReplacement(container_id="comment:999:p:0", text="nope"),
            ],
            strict=True,
        )

    assert list(doc.texts) == original_texts
    assert [comment.text for comment in doc.comments] == original_comments
    assert "CHANGED" not in ZipFile(path).read("word/document.xml").decode("utf-8")


def test_invalid_comment_offsets_fail_before_an_earlier_target_is_applied(
    tmp_path: Path,
) -> None:
    path = _plain_comment_docx(tmp_path / "bad-offset.docx")
    doc = DocxDocument.open(path)
    original = list(doc.texts)

    with pytest.raises(ValueError, match="invalid replacement offsets"):
        doc.apply_replacements(
            [
                SegmentReplacement(container_id="body:p:0", text="CHANGED"),
                SegmentReplacement(
                    container_id="comment:0:p:0",
                    text="nope",
                    start_offset=0,
                    end_offset=999,
                ),
            ],
            strict=True,
        )

    assert list(doc.texts) == original


def test_move_from_inside_comment_fails_closed_without_partial_write(
    tmp_path: Path,
) -> None:
    path = _plain_comment_docx(tmp_path / "move-comment.docx")
    with ZipFile(path) as bundle:
        comments_xml = bundle.read("word/comments.xml")
    comments_root = ElementTree.fromstring(comments_xml)
    comment = comments_root.find(f"{_W}comment")
    assert comment is not None
    move = ElementTree.SubElement(comment, f"{_W}moveFrom")
    move.set(f"{_W}id", "99")
    move.set(f"{_W}author", "Mover")
    move_r = ElementTree.SubElement(move, f"{_W}r")
    move_t = ElementTree.SubElement(move_r, f"{_W}t")
    move_t.text = "migrated"
    _zip_replace(
        path,
        {
            "word/comments.xml": ElementTree.tostring(
                comments_root, encoding="UTF-8", xml_declaration=True
            )
        },
    )

    doc = DocxDocument.open(path)
    original = list(doc.texts)
    with pytest.raises(UnsupportedRevisionError, match="moveFrom"):
        doc.apply_replacements(
            [SegmentReplacement(container_id="body:p:0", text="CHANGED")],
            strict=True,
        )
    assert list(doc.texts) == original



def test_comment_reference_recovers_locator_when_range_start_is_outside_paragraph(
    tmp_path: Path,
) -> None:
    path = _plain_comment_docx(tmp_path / "sdt-start.docx", body="Anchored.")
    with ZipFile(path) as bundle:
        document_xml = bundle.read("word/document.xml")
    root = ElementTree.fromstring(document_xml)
    start = root.find(f".//{_W}commentRangeStart")
    assert start is not None
    body = root.find(f"{_W}body")
    assert body is not None
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    paragraph.remove(start)
    sdt = ElementTree.Element(f"{_W}sdt")
    content = ElementTree.SubElement(sdt, f"{_W}sdtContent")
    content.append(start)
    body.insert(0, sdt)
    _zip_replace(
        path,
        {
            "word/document.xml": ElementTree.tostring(
                root, encoding="UTF-8", xml_declaration=True
            )
        },
    )

    comments = DocxDocument.open(path).comments
    assert len(comments) == 1
    assert comments[0].locator == "body:p:0"


def test_comment_anchor_text_skips_deleted_runs(tmp_path: Path) -> None:
    path = _plain_comment_docx(tmp_path / "visible-anchor.docx", body="Alpha ")
    with ZipFile(path) as bundle:
        document_xml = bundle.read("word/document.xml")
    root = ElementTree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    end = paragraph.find(f"{_W}commentRangeEnd")
    assert end is not None
    end_index = list(paragraph).index(end)
    insertion = ElementTree.Element(f"{_W}ins")
    insertion.set(f"{_W}id", "1")
    insertion.set(f"{_W}author", "Source")
    insert_run = ElementTree.SubElement(insertion, f"{_W}r")
    insert_text = ElementTree.SubElement(insert_run, f"{_W}t")
    insert_text.text = "replacement"
    deletion = ElementTree.Element(f"{_W}del")
    deletion.set(f"{_W}id", "2")
    deletion.set(f"{_W}author", "Source")
    delete_run = ElementTree.SubElement(deletion, f"{_W}r")
    delete_text = ElementTree.SubElement(delete_run, f"{_W}delText")
    delete_text.text = "target"
    paragraph.insert(end_index, insertion)
    paragraph.insert(end_index + 1, deletion)
    _zip_replace(
        path,
        {
            "word/document.xml": ElementTree.tostring(
                root, encoding="UTF-8", xml_declaration=True
            )
        },
    )

    comments = DocxDocument.open(path).comments
    assert len(comments) == 1
    assert comments[0].locator == "body:p:0"
    assert comments[0].anchor_text == "Alpha replacement"
