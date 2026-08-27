from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document as PyDocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from docxtor import (
    AddressableSpan,
    DocxDocument,
    SegmentReplacement,
    UnsupportedRevisionError,
)


def _append_run(parent: object, text: str, *, deleted: bool = False, bold: bool = False) -> None:
    run = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        run.append(rpr)
    node = OxmlElement("w:delText" if deleted else "w:t")
    if text[:1].isspace() or text[-1:].isspace():
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    parent.append(run)  # type: ignore[union-attr]


def _append_ins(
    paragraph: Paragraph,
    text: str,
    *,
    rev_id: str = "10",
    author: str = "Ann Reviewer",
    date: str = "2024-02-03T10:11:12Z",
    bold: bool = False,
) -> None:
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), rev_id)
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    _append_run(ins, text, bold=bold)
    paragraph._p.append(ins)


def _append_del(
    paragraph: Paragraph,
    text: str,
    *,
    rev_id: str = "11",
    author: str = "Bob Reviewer",
    date: str = "2024-02-04T10:11:12Z",
) -> None:
    deletion = OxmlElement("w:del")
    deletion.set(qn("w:id"), rev_id)
    deletion.set(qn("w:author"), author)
    deletion.set(qn("w:date"), date)
    _append_run(deletion, text, deleted=True)
    paragraph._p.append(deletion)


def _append_hyperlink(
    paragraph: Paragraph,
    text: str,
    *,
    anchor: str | None = None,
    rel_id: str | None = None,
    history: str | None = "1",
) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    if anchor is not None:
        hyperlink.set(qn("w:anchor"), anchor)
    if rel_id is not None:
        hyperlink.set(qn("r:id"), rel_id)
    if history is not None:
        hyperlink.set(qn("w:history"), history)
    _append_run(hyperlink, text)
    paragraph._p.append(hyperlink)


def _write_parts(
    path: Path,
    document_xml: str,
    *,
    document_rels: str | None = None,
) -> None:
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    office_rel = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
    )
    main_ct = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
    )
    rels_ct = "application/vnd.openxmlformats-package.relationships+xml"
    with ZipFile(path, mode="w", compression=ZIP_DEFLATED) as archive:
        overrides = (
            f'<Override PartName="/word/document.xml" ContentType="{main_ct}"/>'
        )
        archive.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<Types xmlns="{ct_ns}">'
                f'<Default Extension="rels" ContentType="{rels_ct}"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                f"{overrides}"
                "</Types>"
            ),
        )
        archive.writestr(
            "_rels/.rels",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<Relationships xmlns="{rel_ns}">'
                f'<Relationship Id="rId1" Type="{office_rel}" Target="word/document.xml"/>'
                "</Relationships>"
            ),
        )
        archive.writestr("word/document.xml", document_xml)
        if document_rels is not None:
            archive.writestr("word/_rels/document.xml.rels", document_rels)


def _body_fixture(path: Path) -> str:
    source = PyDocxDocument()
    paragraph = source.add_paragraph()
    paragraph.add_run("See ")
    _append_ins(paragraph, "Alice", bold=True)
    paragraph.add_run(" and ")
    rel_id = paragraph.part.relate_to(
        "https://example.com/contract", RT.HYPERLINK, is_external=True
    )
    _append_hyperlink(paragraph, "contract", rel_id=rel_id)
    paragraph.add_run(" was ")
    _append_del(paragraph, "removed")
    source.save(str(path))
    return rel_id


def _span(doc: DocxDocument, role: str, text: str) -> AddressableSpan:
    matches = [span for span in doc.spans if span.role == role and span.text == text]
    assert matches, f"missing {role} span {text!r} in {list(doc.spans)}"
    return matches[0]


def test_extracts_ins_del_and_hyperlink_spans_with_stable_ids(tmp_path: Path) -> None:
    path = tmp_path / "spans.docx"
    rel_id = _body_fixture(path)

    doc = DocxDocument.open(path)

    assert doc.texts[0] == "See Alice and contract was removed"
    roles = [(span.role, span.text, span.span_id) for span in doc.spans]
    assert roles == [
        ("run", "See ", "body:p:0:span:0"),
        ("insertion", "Alice", "body:p:0:span:1"),
        ("run", " and ", "body:p:0:span:2"),
        ("hyperlink", "contract", "body:p:0:span:3"),
        ("run", " was ", "body:p:0:span:4"),
        ("deletion", "removed", "body:p:0:span:5"),
    ]

    inserted = _span(doc, "insertion", "Alice")
    assert inserted.container_id == "body:p:0"
    assert inserted.start_offset == 4
    assert inserted.end_offset == 9
    assert inserted.revision_id == "10"
    assert inserted.revision_author == "Ann Reviewer"
    assert inserted.revision_date == "2024-02-03T10:11:12Z"

    deleted = _span(doc, "deletion", "removed")
    assert deleted.revision_id == "11"
    assert deleted.revision_author == "Bob Reviewer"

    link = _span(doc, "hyperlink", "contract")
    assert link.hyperlink_rel_id == rel_id
    assert link.hyperlink_anchor is None


def test_partial_and_span_replacements_preserve_wrappers(tmp_path: Path) -> None:
    path = tmp_path / "edit.docx"
    out = tmp_path / "edit-out.docx"
    rel_id = _body_fixture(path)
    doc = DocxDocument.open(path)

    doc.apply_replacements(
        [
            SegmentReplacement(span_id="body:p:0:span:1", text="Alina"),
            SegmentReplacement(
                span_id="body:p:0:span:3",
                text="X",
                start_offset=0,
                end_offset=3,
            ),
            SegmentReplacement(
                container_id="body:p:0",
                text="hidden",
                start_offset=len("See Alina and Xtract was "),
                end_offset=len("See Alina and Xtract was removed"),
            ),
        ],
        strict=True,
    )
    doc.save_docx(out)

    reopened = DocxDocument.open(out)
    assert reopened.texts[0] == "See Alina and Xtract was hidden"
    inserted = _span(reopened, "insertion", "Alina")
    assert inserted.revision_id == "10"
    assert inserted.revision_author == "Ann Reviewer"
    assert inserted.revision_date == "2024-02-03T10:11:12Z"
    link = _span(reopened, "hyperlink", "Xtract")
    assert link.hyperlink_rel_id == rel_id
    deleted = _span(reopened, "deletion", "hidden")
    assert deleted.revision_id == "11"
    assert deleted.revision_author == "Bob Reviewer"

    xml = ZipFile(out).read("word/document.xml").decode("utf-8")
    assert 'w:author="Ann Reviewer"' in xml
    assert f'r:id="{rel_id}"' in xml
    assert "<w:ins" in xml and "</w:ins>" in xml
    assert "<w:del" in xml and "</w:del>" in xml
    assert "<w:hyperlink" in xml and "</w:hyperlink>" in xml
    assert "<w:delText" in xml
    assert "Alina" in xml and "Alice" not in xml


def test_toc_hyperlink_keeps_anchor_after_display_text_replace(tmp_path: Path) -> None:
    path = tmp_path / "toc.docx"
    out = tmp_path / "toc-out.docx"
    source = PyDocxDocument()
    paragraph = source.add_paragraph()
    _append_hyperlink(paragraph, "1. Scope of processing", anchor="_Toc123456", history="1")
    source.save(str(path))

    doc = DocxDocument.open(path)
    link = _span(doc, "hyperlink", "1. Scope of processing")
    assert link.hyperlink_anchor == "_Toc123456"
    doc.apply_replacements(
        [SegmentReplacement(span_id=link.span_id, text="1. Zakres przetwarzania")],
        strict=True,
    )
    doc.save_docx(out)

    reopened = DocxDocument.open(out)
    updated = _span(reopened, "hyperlink", "1. Zakres przetwarzania")
    assert updated.hyperlink_anchor == "_Toc123456"
    assert updated.hyperlink_rel_id is None
    xml = ZipFile(out).read("word/document.xml").decode("utf-8")
    assert 'w:anchor="_Toc123456"' in xml
    assert "1. Scope of processing" not in xml


def test_nested_markup_across_indexed_story_parts(tmp_path: Path) -> None:
    path = tmp_path / "stories.docx"
    out = tmp_path / "stories-out.docx"
    source = PyDocxDocument()
    body = source.add_paragraph()
    body.add_run("Body ")
    _append_ins(body, "body-ins", rev_id="21")

    table = source.add_table(rows=1, cols=1)
    cell_para = table.cell(0, 0).paragraphs[0]
    cell_para.add_run("Cell ")
    _append_hyperlink(cell_para, "cell-link", anchor="_TocCell")

    header_para = source.sections[0].header.paragraphs[0]
    header_para.add_run("Hdr ")
    _append_del(header_para, "hdr-del", rev_id="22")

    footer_para = source.sections[0].footer.paragraphs[0]
    footer_para.add_run("Ftr ")
    _append_ins(footer_para, "ftr-ins", rev_id="23")
    source.save(str(path))

    box_path = tmp_path / "box.docx"
    _write_parts(
        box_path,
        """
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    <w:p><w:r><w:t>Outer.</w:t></w:r></w:p>
    <w:p>
      <w:r>
        <w:pict xmlns:v="urn:schemas-microsoft-com:vml">
          <v:shape>
            <v:textbox>
              <w:txbxContent>
                <w:p>
                  <w:r><w:t>Box </w:t></w:r>
                  <w:hyperlink w:anchor="_TocBox" w:history="1">
                    <w:r><w:t>box-link</w:t></w:r>
                  </w:hyperlink>
                </w:p>
              </w:txbxContent>
            </v:textbox>
          </v:shape>
        </w:pict>
      </w:r>
    </w:p>
  </w:body>
</w:document>
""".strip(),
    )

    doc = DocxDocument.open(path)
    by_container = {span.container_id: span for span in doc.spans if span.role != "run"}
    assert by_container["body:p:0"].role == "insertion"
    assert by_container["body:p:0"].text == "body-ins"
    table_span = next(
        span
        for span in doc.spans
        if span.container_id.startswith("table:") and span.role == "hyperlink"
    )
    assert table_span.hyperlink_anchor == "_TocCell"
    header_span = next(
        span
        for span in doc.spans
        if span.container_id.startswith("header:") and span.role == "deletion"
    )
    footer_span = next(
        span
        for span in doc.spans
        if span.container_id.startswith("footer:") and span.role == "insertion"
    )

    doc.apply_replacements(
        [
            SegmentReplacement(span_id=by_container["body:p:0"].span_id, text="BODY"),
            SegmentReplacement(span_id=table_span.span_id, text="CELL"),
            SegmentReplacement(span_id=header_span.span_id, text="HDR"),
            SegmentReplacement(span_id=footer_span.span_id, text="FTR"),
        ],
        strict=True,
    )
    doc.save_docx(out)
    roundtrip = DocxDocument.open(out)
    assert "BODY" in roundtrip.texts[0]
    assert any(
        span.text == "CELL" and span.hyperlink_anchor == "_TocCell"
        for span in roundtrip.spans
    )
    assert any(span.text == "HDR" and span.role == "deletion" for span in roundtrip.spans)
    assert any(span.text == "FTR" and span.role == "insertion" for span in roundtrip.spans)

    box = DocxDocument.open(box_path)
    box_span = next(
        span
        for span in box.spans
        if span.container_id.startswith("txbx:") and span.role == "hyperlink"
    )
    assert box_span.role == "hyperlink"
    assert box_span.text == "box-link"
    box.apply_replacements(
        [SegmentReplacement(span_id=box_span.span_id, text="BOX")],
        strict=True,
    )
    box_out = tmp_path / "box-out.docx"
    box.save_docx(box_out)
    box_rt = DocxDocument.open(box_out)
    updated_box = next(
        span
        for span in box_rt.spans
        if span.container_id.startswith("txbx:") and span.role == "hyperlink"
    )
    assert updated_box.text == "BOX"
    assert updated_box.hyperlink_anchor == "_TocBox"


def test_untouched_document_keeps_revision_and_hyperlink_wrappers(tmp_path: Path) -> None:
    path = tmp_path / "untouched.docx"
    out = tmp_path / "untouched-out.docx"
    rel_id = _body_fixture(path)
    original = ZipFile(path).read("word/document.xml").decode("utf-8")

    doc = DocxDocument.open(path)
    assert [span.role for span in doc.spans] == [
        "run",
        "insertion",
        "run",
        "hyperlink",
        "run",
        "deletion",
    ]
    doc.save_docx(out)

    xml = ZipFile(out).read("word/document.xml").decode("utf-8")
    assert "<w:ins" in xml and 'w:author="Ann Reviewer"' in xml
    assert "<w:del" in xml and "<w:delText" in xml
    assert "<w:hyperlink" in xml and f'r:id="{rel_id}"' in xml
    assert "Alice" in xml and "contract" in xml and "removed" in xml
    reopened = DocxDocument.open(out)
    assert reopened.texts[0] == "See Alice and contract was removed"
    assert _span(reopened, "hyperlink", "contract").hyperlink_rel_id == rel_id
    assert "w:ins" in original


def test_unsupported_move_from_fails_closed_without_partial_write(tmp_path: Path) -> None:
    path = tmp_path / "move.docx"
    source = PyDocxDocument()
    first = source.add_paragraph("Keep me")
    first.add_run("")
    second = source.add_paragraph()
    move = OxmlElement("w:moveFrom")
    move.set(qn("w:id"), "99")
    move.set(qn("w:author"), "Mover")
    _append_run(move, "migrated")
    second._p.append(move)
    source.save(str(path))

    doc = DocxDocument.open(path)
    original = list(doc.texts)
    with pytest.raises(UnsupportedRevisionError, match="moveFrom"):
        doc.apply_replacements(
            [SegmentReplacement(container_id="body:p:0", text="CHANGED")],
            strict=True,
        )
    assert list(doc.texts) == original
    assert "Keep me" in ZipFile(path).read("word/document.xml").decode("utf-8")
    assert "CHANGED" not in ZipFile(path).read("word/document.xml").decode("utf-8")


def test_unsupported_block_level_ins_fails_closed(tmp_path: Path) -> None:
    path = tmp_path / "block-ins.docx"
    _write_parts(
        path,
        """
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>Visible.</w:t></w:r></w:p>
    <w:ins w:id="7" w:author="Ann">
      <w:p><w:r><w:t>Inserted paragraph.</w:t></w:r></w:p>
    </w:ins>
  </w:body>
</w:document>
""".strip(),
    )

    doc = DocxDocument.open(path)
    original = list(doc.texts)
    with pytest.raises(UnsupportedRevisionError, match="block-ins"):
        doc.apply_replacements(
            [SegmentReplacement(container_id="body:p:0", text="NEW")],
            strict=True,
        )
    assert list(doc.texts) == original


def test_nested_hyperlink_inside_insertion_keeps_both_wrappers(tmp_path: Path) -> None:
    path = tmp_path / "nested.docx"
    out = tmp_path / "nested-out.docx"
    source = PyDocxDocument()
    paragraph = source.add_paragraph()
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "3")
    ins.set(qn("w:author"), "Ann Reviewer")
    ins.set(qn("w:date"), "2024-05-06T00:00:00Z")
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "_TocNested")
    _append_run(hyperlink, "inner-link")
    ins.append(hyperlink)
    paragraph._p.append(ins)
    source.save(str(path))

    doc = DocxDocument.open(path)
    span = doc.spans[0]
    assert span.role == "hyperlink"
    assert span.text == "inner-link"
    assert span.revision_id == "3"
    assert span.revision_author == "Ann Reviewer"
    assert span.hyperlink_anchor == "_TocNested"

    doc.apply_replacements(
        [SegmentReplacement(span_id=span.span_id, text="INNER")],
        strict=True,
    )
    doc.save_docx(out)
    xml = ZipFile(out).read("word/document.xml").decode("utf-8")
    assert xml.index("<w:ins") < xml.index("<w:hyperlink")
    assert 'w:anchor="_TocNested"' in xml
    assert 'w:author="Ann Reviewer"' in xml
    reopened = DocxDocument.open(out)
    updated = reopened.spans[0]
    assert updated.text == "INNER"
    assert updated.role == "hyperlink"
    assert updated.revision_id == "3"
    assert updated.hyperlink_anchor == "_TocNested"
