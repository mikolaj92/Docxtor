from __future__ import annotations

from dataclasses import dataclass
from typing import Any

from docx.document import Document as DocxDocumentType
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .docx_comments import (
    _capture_thread_parts,
    _collect_comments,
    _ensure_thread_parts,
    _existing_comments_part,
)
from .docx_models import AddressableComment, TextSegment
from .docx_ns import W_P, W_SDT, W_SDT_CONTENT, W_TBL
from .docx_units import _paragraph_visible_text
from .docx_xml import _is_text_box_container, _local_tag, _w_get


def _existing_note_part(doc: DocxDocumentType, reltype: str) -> tuple[Any, Any] | None:
    try:
        part = doc.part.part_related_by(reltype)
    except KeyError:
        return None
    return part, parse_xml(part.blob)


def _iter_text_box_hosts(root: Any) -> list[Any]:
    """Return each text-bearing ``w:txbxContent`` (or VML/DrawingML wrapper).

    Nested boxes keep document order. Empty decorative shapes (no descendant
    ``w:t``) are skipped so they never occupy a ``txbx:N`` slot.
    """
    hosts: list[Any] = []
    seen: set[int] = set()

    def has_text(element: Any) -> bool:
        for node in element.iter():
            if _local_tag(node.tag) in {"t", "delText"} and node.text and node.text.strip():
                return True
        return False

    for element in root.iter():
        if not _is_text_box_container(element.tag):
            continue
        identity = id(element)
        if identity in seen:
            continue
        if _local_tag(element.tag) != "txbxContent":
            inner = None
            for child in element.iter():
                if child is not element and _local_tag(child.tag) == "txbxContent":
                    inner = child
                    break
            if inner is not None:
                if id(inner) in seen:
                    continue
                if not has_text(inner):
                    seen.add(id(inner))
                    continue
                seen.add(id(inner))
                hosts.append(inner)
                continue
        if not has_text(element):
            seen.add(identity)
            continue
        seen.add(identity)
        hosts.append(element)
    return hosts


def _iter_paragraph_elements(container: Any, *, skip_text_boxes: bool = False) -> list[Any]:
    """Collect w:p elements in document order, including those nested in w:sdt.

    python-docx's ``.paragraphs`` only returns direct ``w:p`` children and therefore
    omits content-control (``w:sdt`` / ``w:sdtContent``) paragraphs. Tables are left
    to the caller's existing table walk so global ordering stays unchanged.

    Floating text boxes (``w:txbxContent``) are a separate container family. The
    body/header/table walk skips them so boxed paragraphs are not double-counted
    as body runs; ``_iter_text_box_hosts`` indexes them as ``txbx:N``.
    """
    result: list[Any] = []

    def walk(element: Any) -> None:
        for child in element:
            tag = child.tag
            if skip_text_boxes and _is_text_box_container(tag):
                continue
            if tag == W_P:
                result.append(child)
            elif tag == W_SDT:
                for content in child.iterchildren(W_SDT_CONTENT):
                    walk(content)
            # Nested tables are handled by the dedicated table enumeration path.

    walk(container)
    return result


def _paragraphs_from_container(
    container_element: Any,
    parent: Any,
    *,
    skip_text_boxes: bool = False,
) -> list[Paragraph]:
    """Wrap collected paragraph elements as python-docx Paragraph proxies."""
    return [
        Paragraph(p, parent)
        for p in _iter_paragraph_elements(container_element, skip_text_boxes=skip_text_boxes)
    ]


@dataclass
class _ParaRef:
    """Internal mapping from our segment to python-docx paragraph + metadata."""

    id: str
    container_id: str
    paragraph_index: int | None
    paragraph: Paragraph
    part_name: str  # "body", "header:0", "table:0:r:0:c:0", etc.


@dataclass
class IndexedStories:
    """Mechanical walk of one python-docx document into addressable stories."""

    segments: list[TextSegment]
    refs: list[_ParaRef]
    paragraphs_by_index: dict[int, Paragraph]
    paragraphs_by_container: dict[str, Paragraph]
    comments: list[AddressableComment]
    thread_parts: dict[str, tuple[bytes, str]]
    note_parts: dict[str, tuple[Any, Any]]


def index_stories(doc: DocxDocumentType) -> IndexedStories:
    """Walk body, tables, headers, text boxes, comments, and notes."""
    segments: list[TextSegment] = []
    refs: list[_ParaRef] = []

    # Global paragraph index counts EVERY paragraph in document order
    # (body, table cells, headers, footers, text boxes), including empty
    # ones. This matches the contract expected by dike_docs locator and
    # anchors. Text boxes are appended after headers/footers so documents
    # without boxes keep existing body/table/header indices.
    global_paragraph_index = 0
    body_paragraph_index = 0
    paragraphs_by_index: dict[int, Paragraph] = {}
    paragraphs_by_container: dict[str, Paragraph] = {}

    def add_paragraphs(paragraphs: list[Paragraph], prefix: str) -> None:
        nonlocal body_paragraph_index, global_paragraph_index
        for local_idx, para in enumerate(paragraphs):
            paragraphs_by_index[global_paragraph_index] = para

            # container_id: body uses global index for stability (matches Dike anchors);
            # other sections use local index within their container.
            if prefix == "body":
                cid = f"body:p:{body_paragraph_index}"
                body_paragraph_index += 1
            else:
                cid = f"{prefix}:p:{local_idx}"

            paragraphs_by_container[cid] = para

            text = _paragraph_visible_text(para)
            run_indices = [ri for ri, run in enumerate(para.runs) if run.text] if para.runs else []

            if text:
                seg_id = f"s{len(segments)}"
                segments.append(
                    TextSegment(
                        id=seg_id,
                        text=text,
                        part=(
                            "word/comments.xml"
                            if prefix.startswith("comment:")
                            else "word/document.xml"
                            if prefix.startswith(("body", "table", "txbx"))
                            else f"word/{prefix.split(':')[0]}.xml"
                        ),
                        index=local_idx,
                        container_id=cid,
                        paragraph_index=global_paragraph_index,
                        run_indices=run_indices,
                    )
                )
                refs.append(
                    _ParaRef(
                        id=seg_id,
                        container_id=cid,
                        paragraph_index=global_paragraph_index,
                        paragraph=para,
                        part_name=prefix,
                    )
                )

            global_paragraph_index += 1

    def add_comment_paragraphs(paragraphs: list[Paragraph], comment_id: str) -> None:
        for local_idx, para in enumerate(paragraphs):
            cid = f"comment:{comment_id}:p:{local_idx}"
            paragraphs_by_container[cid] = para
            text = _paragraph_visible_text(para)
            if not text:
                continue
            seg_id = f"s{len(segments)}"
            segments.append(
                TextSegment(
                    id=seg_id,
                    text=text,
                    part="word/comments.xml",
                    index=local_idx,
                    container_id=cid,
                    paragraph_index=None,
                    run_indices=[run_index for run_index, run in enumerate(para.runs) if run.text],
                )
            )
            refs.append(
                _ParaRef(
                    id=seg_id,
                    container_id=cid,
                    paragraph_index=None,
                    paragraph=para,
                    part_name=f"comment:{comment_id}",
                )
            )

    def add_note_paragraphs(
        paragraphs: list[Paragraph], story: str, note_id: str, part_name: str
    ) -> None:
        for local_idx, para in enumerate(paragraphs):
            cid = f"{story}:{note_id}:p:{local_idx}"
            paragraphs_by_container[cid] = para
            text = _paragraph_visible_text(para)
            if not text:
                continue
            seg_id = f"s{len(segments)}"
            segments.append(
                TextSegment(
                    id=seg_id,
                    text=text,
                    part=part_name,
                    index=local_idx,
                    container_id=cid,
                    paragraph_index=None,
                    run_indices=[run_index for run_index, run in enumerate(para.runs) if run.text],
                )
            )
            refs.append(
                _ParaRef(
                    id=seg_id,
                    container_id=cid,
                    paragraph_index=None,
                    paragraph=para,
                    part_name=f"{story}:{note_id}",
                )
            )

    # Body paragraphs and tables in authored XML order. Docxtor owns this
    # mechanical block walk so consumers never rebuild a second OOXML order map.
    # Body ids still count body paragraphs only; table ids still use python-docx's
    # table/row/cell coordinates and unique physical cells (#36 / #48).
    table_index = 0
    for block in doc.element.body.iterchildren():
        if block.tag == W_P:
            add_paragraphs([Paragraph(block, doc._body)], "body")
            continue
        if block.tag == W_SDT:
            content = block.find(W_SDT_CONTENT)
            paragraphs = (
                []
                if content is None
                else _paragraphs_from_container(content, doc._body, skip_text_boxes=True)
            )
            add_paragraphs(paragraphs, "body")
            continue
        if block.tag != W_TBL:
            continue
        table = doc.tables[table_index]
        seen_cells: set[object] = set()
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                add_paragraphs(
                    _paragraphs_from_container(cell._tc, cell, skip_text_boxes=True),
                    f"table:{table_index}:r:{ri}:c:{ci}",
                )
        table_index += 1

    # Headers / Footers. Accessing ``._element`` on an absent or linked
    # story calls python-docx's get-or-add path and mutates the package.
    # Only inspect definitions explicitly referenced by this section.
    for si, section in enumerate(doc.sections):
        for story_name, story in (
            ("header", section.header),
            ("header-first", section.first_page_header),
            ("header-even", section.even_page_header),
            ("footer", section.footer),
            ("footer-first", section.first_page_footer),
            ("footer-even", section.even_page_footer),
        ):
            if not story._has_definition:
                continue
            definition = story._definition
            add_paragraphs(
                _paragraphs_from_container(definition.element, story, skip_text_boxes=True),
                f"{story_name}:{si}",
            )

    # Floating text boxes (VML v:textbox / w:txbxContent / DrawingML wps:txbx).
    # python-docx does not surface these as paragraphs. Index after the
    # ordinary stories so documents without boxes keep stable body ids.
    for box_idx, host in enumerate(_iter_text_box_hosts(doc.element)):
        add_paragraphs(
            _paragraphs_from_container(host, doc._body, skip_text_boxes=True),
            f"txbx:{box_idx}",
        )

    comments_part = _existing_comments_part(doc)
    if comments_part is not None:
        for comment_elm in comments_part.element.findall(qn("w:comment")):
            comment_id = _w_get(comment_elm, "id")
            if comment_id is None:
                continue
            add_comment_paragraphs(
                _paragraphs_from_container(comment_elm, comments_part),
                comment_id,
            )

    note_parts: dict[str, tuple[Any, Any]] = {}
    for story, reltype, root_tag, note_tag, part_name in (
        ("footnote", RT.FOOTNOTES, "w:footnotes", "w:footnote", "word/footnotes.xml"),
        ("endnote", RT.ENDNOTES, "w:endnotes", "w:endnote", "word/endnotes.xml"),
    ):
        existing_note = _existing_note_part(doc, reltype)
        if existing_note is None:
            continue
        note_part, note_root = existing_note
        if note_root.tag != qn(root_tag):
            continue
        note_parts[story] = (note_part, note_root)
        for note_elm in note_root.findall(qn(note_tag)):
            note_id = _w_get(note_elm, "id")
            note_type = _w_get(note_elm, "type")
            if note_id is None or note_type in {"separator", "continuationSeparator"}:
                continue
            add_note_paragraphs(
                _paragraphs_from_container(note_elm, note_part),
                story,
                note_id,
                part_name,
            )

    comments = _collect_comments(comments_part, paragraphs_by_container, doc.part.package)
    thread_parts = _capture_thread_parts(doc.part.package)
    _ensure_thread_parts(doc.part.package, thread_parts)
    return IndexedStories(
        segments=segments,
        refs=refs,
        paragraphs_by_index=paragraphs_by_index,
        paragraphs_by_container=paragraphs_by_container,
        comments=comments,
        thread_parts=thread_parts,
        note_parts=note_parts,
    )
