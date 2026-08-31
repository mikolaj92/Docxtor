from __future__ import annotations

import os
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.document import Document as PythonDocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .docx_package import normalize_docx_timestamps


class PublicationMarkError(ValueError):
    """A physical DOCX publication mark could not be applied safely."""


@dataclass(frozen=True)
class DocumentMark:
    content_status: str
    subject: str
    keywords: tuple[str, ...]
    footer_text: str


@dataclass(frozen=True)
class BodyAppendix:
    heading: str
    paragraphs: tuple[str, ...]


def has_document_mark(source: str | Path | bytes, mark: DocumentMark) -> bool:
    document = _open(source)
    status = (document.core_properties.content_status or "").strip()
    if status != mark.content_status:
        return False
    return any(
        mark.footer_text in paragraph.text
        for footer in _footers(document)
        for paragraph in footer.paragraphs
    )


def stamp_document_mark(source: str | Path | bytes, mark: DocumentMark) -> bytes:
    document = _open(source)
    props = document.core_properties
    props.content_status = mark.content_status
    props.subject = mark.subject
    existing = tuple(item.strip() for item in (props.keywords or "").split(";") if item.strip())
    props.keywords = ";".join(dict.fromkeys((*existing, *mark.keywords)))
    for footer in _footers(document):
        if footer.is_linked_to_previous:
            footer.is_linked_to_previous = False
        if any(mark.footer_text in paragraph.text for paragraph in footer.paragraphs):
            continue
        if footer.paragraphs and not footer.paragraphs[0].text.strip():
            footer.paragraphs[0].text = mark.footer_text
        else:
            footer.add_paragraph(mark.footer_text)
    data = _save(document)
    if not has_document_mark(data, mark):
        raise PublicationMarkError("DOCX did not retain the requested document mark")
    return data


def has_body_appendix(source: str | Path | bytes, *, heading: str) -> bool:
    document = _open(source)
    return _heading_index(document, heading) is not None


def append_body_appendix(source: str | Path | bytes, appendix: BodyAppendix) -> bytes:
    document = _open(source)
    if _heading_index(document, appendix.heading) is not None:
        raise PublicationMarkError("DOCX already carries the requested body appendix")
    document.add_paragraph("")
    heading = document.add_paragraph()
    heading.add_run(appendix.heading).bold = True
    for text in appendix.paragraphs:
        document.add_paragraph(text)
    data = _save(document)
    if not has_body_appendix(data, heading=appendix.heading):
        raise PublicationMarkError("DOCX did not retain the requested body appendix")
    return data


def remove_body_appendix(source: str | Path | bytes, *, heading: str) -> tuple[bytes, bool]:
    document = _open(source)
    index = _heading_index(document, heading)
    if index is None:
        return _source_bytes(source), False
    start = index
    if start > 0 and not document.paragraphs[start - 1].text.strip():
        start -= 1
    body = document.element.body
    for paragraph in list(document.paragraphs[start:]):
        element = paragraph._element
        if element.getparent() is body:
            body.remove(element)
    data = _save(document)
    if has_body_appendix(data, heading=heading):
        raise PublicationMarkError("DOCX retained the removed body appendix")
    return data, True


def write_publication_bytes(path: str | Path, data: bytes) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            dir=target.parent, prefix=f".{target.name}.", suffix=".tmp", delete=False
        ) as handle:
            temporary = Path(handle.name)
            handle.write(data)
        normalize_docx_timestamps(temporary)
        os.replace(temporary, target)
        temporary = None
        return target
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def _source_bytes(source: str | Path | bytes) -> bytes:
    return source if isinstance(source, bytes) else Path(source).read_bytes()


def _open(source: str | Path | bytes) -> PythonDocxDocument:
    try:
        return Document(BytesIO(_source_bytes(source)))
    except (OSError, ValueError, KeyError, BadZipFile, PackageNotFoundError) as exc:
        raise PublicationMarkError(f"DOCX publication package is unreadable: {exc}") from exc


def _save(document: PythonDocxDocument) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _heading_index(document: PythonDocxDocument, heading: str) -> int | None:
    return next(
        (i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == heading),
        None,
    )


def _footers(document: PythonDocxDocument) -> tuple[object, ...]:
    result: list[object] = []
    for section in document.sections:
        candidates = [section.footer]
        if section.different_first_page_header_footer:
            candidates.append(section.first_page_footer)
        if document.settings.odd_and_even_pages_header_footer:
            candidates.append(section.even_page_footer)
        result.extend(candidates)
    return tuple(result)
