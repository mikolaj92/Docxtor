from __future__ import annotations

from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
from typing import Any, cast

from docx import Document as PyDocxDocument
from docx.oxml.ns import qn

from .common import DocumentError
from .docx_inline import (
    _index_at_visible_offset,
    _split_visible_offset,
    _visible_text,
    paragraph_to_inline_segments,
    rebuild_paragraph_from_inline,
)
from .docx_models import AddressableComment
from .docx_review_models import OperationReceipt, OperationStatus
from .docx_stories import index_stories


class CommentMutationError(DocumentError):
    """A comment mutation could not be proven complete."""


@dataclass(frozen=True)
class CommentRange:
    locator: str
    start_offset: int
    end_offset: int
    expected_text: str | None = None


@dataclass(frozen=True)
class CommentAuthor:
    author: str
    initials: str | None = None
    date: str | None = None


@dataclass(frozen=True)
class CommentMutationResult:
    data: bytes
    receipt: OperationReceipt
    comments: tuple[AddressableComment, ...]


def add_comment(
    data: bytes,
    target: CommentRange,
    text: str,
    author: CommentAuthor,
) -> CommentMutationResult:
    """Add one comment to an exact range in the canonical story locator space."""
    if not text:
        raise CommentMutationError("comment text must not be empty")
    doc = PyDocxDocument(BytesIO(data))
    stories = index_stories(doc)
    paragraph = stories.paragraphs_by_container.get(target.locator)
    if paragraph is None:
        raise CommentMutationError(f"unknown comment locator: {target.locator}")
    segments = paragraph_to_inline_segments(paragraph)
    visible = _visible_text(segments)
    if not 0 <= target.start_offset < target.end_offset <= len(visible):
        raise CommentMutationError(
            f"invalid comment range {target.start_offset}:{target.end_offset} for {target.locator}"
        )
    selected = visible[target.start_offset : target.end_offset]
    if target.expected_text is not None and selected != target.expected_text:
        raise CommentMutationError(f"comment range text changed at {target.locator}")
    segments = _split_visible_offset(
        _split_visible_offset(segments, target.end_offset), target.start_offset
    )
    start_index = _index_at_visible_offset(segments, target.start_offset)
    end_index = _index_at_visible_offset(segments, target.end_offset) - 1
    if start_index < 0 or end_index < start_index:
        raise CommentMutationError(f"comment range cannot be represented at {target.locator}")
    if any(segment.kind != "text" for segment in segments[start_index : end_index + 1]):
        raise CommentMutationError(f"comment range crosses opaque content at {target.locator}")
    rebuild_paragraph_from_inline(paragraph, segments)
    runs = paragraph.runs
    comment = doc.comments.add_comment(text=text, author=author.author, initials=author.initials)
    if author.date is not None:
        comment._comment_elm.set(qn("w:date"), author.date)
    comment_id = int(comment.comment_id)
    runs[start_index].mark_comment_range(runs[end_index], comment_id)
    # Rebuild places opaque markers verbatim; python-docx then serializes all parts.
    payload = _serialize(doc)
    after = index_stories(PyDocxDocument(BytesIO(payload)))
    created = next((item for item in after.comments if item.comment_id == str(comment_id)), None)
    if created is None or created.locator != target.locator or created.anchor_text != selected:
        raise CommentMutationError("comment creation was not confirmed after round-trip")
    return CommentMutationResult(
        data=payload,
        receipt=OperationReceipt(
            operation="add_comment",
            status=OperationStatus.APPLIED,
            affected_parts=("word/comments.xml",),
            created_ids=(str(comment_id),),
            locator=target.locator,
            before_sha256=sha256(data).hexdigest(),
            after_sha256=sha256(payload).hexdigest(),
        ),
        comments=tuple(after.comments),
    )


def remove_comments(data: bytes, comment_ids: set[str] | None = None) -> CommentMutationResult:
    """Remove selected comments (or every comment) and all range/reference markers."""
    doc = PyDocxDocument(BytesIO(data))
    stories = index_stories(doc)
    existing = {comment.comment_id for comment in stories.comments}
    selected = existing if comment_ids is None else set(comment_ids)
    missing = selected - existing
    if missing:
        raise CommentMutationError(f"unknown comment IDs: {sorted(missing)}")
    for root in _story_roots(doc):
        for node in list(root.iter()):
            if node.tag not in {
                qn("w:commentRangeStart"),
                qn("w:commentRangeEnd"),
                qn("w:commentReference"),
            }:
                continue
            comment_id = node.get(qn("w:id"))
            if comment_id not in selected:
                continue
            owner = node.getparent()
            if (
                owner is not None
                and node.tag == qn("w:commentReference")
                and owner.tag == qn("w:r")
            ):
                parent = owner.getparent()
                if parent is not None:
                    parent.remove(owner)
            elif owner is not None:
                owner.remove(node)
    comments_part = cast(
        Any,
        doc.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        ),
    )
    for comment in list(comments_part.element.findall(qn("w:comment"))):
        if comment.get(qn("w:id")) in selected:
            comments_part.element.remove(comment)
    payload = _serialize(doc)
    after = index_stories(PyDocxDocument(BytesIO(payload)))
    remaining = {comment.comment_id for comment in after.comments}
    if selected & remaining:
        raise CommentMutationError("comment removal was not confirmed after round-trip")
    status = OperationStatus.APPLIED if selected else OperationStatus.NOOP
    return CommentMutationResult(
        data=payload,
        receipt=OperationReceipt(
            operation="remove_comments",
            status=status,
            affected_parts=("word/comments.xml",) if selected else (),
            created_ids=(),
            before_sha256=sha256(data).hexdigest(),
            after_sha256=sha256(payload).hexdigest(),
        ),
        comments=tuple(after.comments),
    )


def _story_roots(doc: Any) -> list[Any]:
    roots = [doc.element]
    for section in doc.sections:
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if story._has_definition:
                roots.append(story._definition.element)
    return roots


def _serialize(doc: Any) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
